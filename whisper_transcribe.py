import whisper
from docx import Document

# Load Whisper model
model = whisper.load_model("large")  # You can use "small", "medium", or "large" for better accuracy

# Path to your audio file
audio_path = r"C:\Users\kalin\Documents\Dict\schwarick.flac"

# Transcribe the audio file
print("Transcribing...")
result = model.transcribe(audio_path, language="de")
text = result["text"]

# Save to a Word document
doc = Document()
doc.add_heading("Transcription", level=1)
doc.add_paragraph(text)
output_path = r"C:\Users\kalin\Documents\schwarick.docx"
doc.save(output_path)

print(f"Transcription saved to {output_path}")
