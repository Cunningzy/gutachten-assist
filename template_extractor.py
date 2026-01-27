"""
Template Extractor - Phases 0-5
Extracts reusable template from example Gutachten documents.

Usage:
    python template_extractor.py extract <folder_with_docx>
    python template_extractor.py analyze <doc_profiles_folder>
"""

import sys
import os
import json
import re
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional
from collections import defaultdict

# Force UTF-8 for Windows
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# =============================================================================
# CONFIGURATION
# =============================================================================

CONFIG = {
    "boilerplate_threshold": 0.85,  # 85% occurrence = fixed block
    "min_text_length": 3,           # Ignore very short paragraphs for boilerplate
    "anchor_similarity_threshold": 0.88,  # Fuzzy matching threshold
    "ngram_sizes": [2, 3, 4, 5],    # Sequence detection lengths
}

# =============================================================================
# PHASE 0: INGEST - Parse DOCX to DocProfile
# =============================================================================

def extract_paragraph_info(para) -> dict:
    """Extract all relevant info from a paragraph."""
    # Get style info
    style_name = para.style.name if para.style else "Normal"

    # Get outline level (heading level)
    outline_level = None
    if style_name.startswith("Heading"):
        try:
            outline_level = int(style_name.replace("Heading ", "").strip())
        except:
            pass

    # Check for list/numbering
    is_list = False
    list_level = None
    numPr = para._p.pPr.numPr if para._p.pPr is not None and para._p.pPr.numPr is not None else None
    if numPr is not None:
        is_list = True
        ilvl = numPr.ilvl
        if ilvl is not None:
            list_level = ilvl.val

    # Get spacing
    spacing_before = None
    spacing_after = None
    if para.paragraph_format.space_before:
        spacing_before = para.paragraph_format.space_before.pt
    if para.paragraph_format.space_after:
        spacing_after = para.paragraph_format.space_after.pt

    # Get indentation
    left_indent = None
    if para.paragraph_format.left_indent:
        left_indent = para.paragraph_format.left_indent.pt

    # Get alignment
    alignment = None
    if para.paragraph_format.alignment:
        alignment = str(para.paragraph_format.alignment)

    return {
        "text": para.text,
        "style": style_name,
        "outline_level": outline_level,
        "is_list": is_list,
        "list_level": list_level,
        "spacing_before": spacing_before,
        "spacing_after": spacing_after,
        "left_indent": left_indent,
        "alignment": alignment,
    }


def extract_header_footer(doc) -> dict:
    """Extract header and footer content."""
    headers = []
    footers = []

    for section in doc.sections:
        # Header
        header = section.header
        if header and header.paragraphs:
            header_text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
            if header_text:
                headers.append(header_text)

        # Footer
        footer = section.footer
        if footer and footer.paragraphs:
            footer_text = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
            if footer_text:
                footers.append(footer_text)

    return {
        "headers": list(set(headers)),  # Deduplicate
        "footers": list(set(footers)),
    }


def extract_styles_used(doc) -> dict:
    """Extract histogram of styles used in document."""
    style_counts = defaultdict(int)
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_counts[style_name] += 1
    return dict(style_counts)


def parse_docx_to_profile(docx_path: str) -> dict:
    """
    Phase 0: Parse a single DOCX file into a DocProfile.
    """
    doc = Document(docx_path)

    paragraphs = []
    for para in doc.paragraphs:
        para_info = extract_paragraph_info(para)
        paragraphs.append(para_info)

    header_footer = extract_header_footer(doc)
    styles_used = extract_styles_used(doc)

    return {
        "source_file": os.path.basename(docx_path),
        "extracted_at": datetime.now().isoformat(),
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "headers": header_footer["headers"],
        "footers": header_footer["footers"],
        "styles_histogram": styles_used,
    }


# =============================================================================
# PHASE 1: NORMALIZE TEXT
# =============================================================================

# Date patterns (German formats)
DATE_PATTERNS = [
    r'\b\d{1,2}\.\d{1,2}\.\d{2,4}\b',      # 27.11.2025, 1.5.25
    r'\b\d{4}-\d{2}-\d{2}\b',               # 2025-11-27
    r'\b\d{1,2}\.\s*(?:Januar|Februar|März|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s*\d{2,4}\b',  # 27. November 2025
    r'\b(?:Januar|Februar|März|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s*\d{2,4}\b',  # November 2025
]

# ID patterns (Aktenzeichen, insurance numbers, etc.)
ID_PATTERNS = [
    r'\b[A-Z]{1,3}\s*\d{6,12}\b',           # Insurance numbers like "A 123456789"
    r'\bAktenzeichen[:\s]*[\w\-/]+\b',      # Aktenzeichen: ABC-123
    r'\bAz\.?[:\s]*[\w\-/]+\b',             # Az.: ABC-123
    r'\bVersicherungsnr?\.?[:\s]*[\w\-/]+\b',  # Versicherungsnr.
    r'\bVers\.?\s*Nr\.?[:\s]*[\w\-/]+\b',   # Vers. Nr.
]

# Name patterns (conservative - only clear patterns)
NAME_PATTERNS = [
    r'\b(?:Herr|Frau|Hr\.|Fr\.)\s+[A-ZÄÖÜ][a-zäöüß]+(?:\s+[A-ZÄÖÜ][a-zäöüß]+)?',  # Herr Müller, Frau Schmidt
    r'\bgeb\.\s*\d{1,2}\.\d{1,2}\.\d{2,4}',  # geb. 15.03.1965
    r'\bgeboren\s+(?:am\s+)?\d{1,2}\.\d{1,2}\.\d{2,4}',  # geboren am 15.03.1965
]


def normalize_basic(text: str) -> str:
    """Basic normalization: lowercase, trim, collapse whitespace, unify punctuation."""
    result = text.lower().strip()
    result = re.sub(r'\s+', ' ', result)  # Collapse whitespace
    result = re.sub(r'[""„]', '"', result)  # Unify quotes
    result = re.sub(r'[''‚]', "'", result)  # Unify apostrophes
    result = re.sub(r'[–—]', '-', result)   # Unify dashes
    return result


def normalize_no_dates(text: str) -> str:
    """Replace dates with <DATE> placeholder."""
    result = text
    for pattern in DATE_PATTERNS:
        result = re.sub(pattern, '<DATE>', result, flags=re.IGNORECASE)
    return result


def normalize_no_ids(text: str) -> str:
    """Replace IDs with <ID> placeholder."""
    result = text
    for pattern in ID_PATTERNS:
        result = re.sub(pattern, '<ID>', result, flags=re.IGNORECASE)
    return result


def normalize_no_names(text: str) -> str:
    """Replace clear name patterns with <NAME> placeholder (conservative)."""
    result = text
    for pattern in NAME_PATTERNS:
        result = re.sub(pattern, '<NAME>', result, flags=re.IGNORECASE)
    return result


def add_normalizations(profile: dict) -> dict:
    """
    Phase 1: Add normalized text variants to each paragraph.
    """
    for para in profile["paragraphs"]:
        text = para["text"]

        # Create normalized variants
        norm_basic = normalize_basic(text)
        norm_no_dates = normalize_no_dates(norm_basic)
        norm_no_ids = normalize_no_ids(norm_no_dates)
        norm_no_names = normalize_no_names(norm_no_ids)  # Most normalized

        para["norm_basic"] = norm_basic
        para["norm_no_dates"] = norm_no_dates
        para["norm_no_ids"] = norm_no_ids
        para["norm_full"] = norm_no_names  # Full normalization

        # Compute fingerprint for boilerplate detection
        fingerprint_text = norm_no_ids  # Use this level for fingerprinting
        fingerprint_hash = hashlib.sha1(fingerprint_text.encode('utf-8')).hexdigest()[:12]
        para["fingerprint"] = f"{para['style']}:{fingerprint_hash}"

    return profile


# =============================================================================
# PHASE 2: BOILERPLATE DETECTION
# =============================================================================

def compute_fingerprint_counts(profiles: list[dict]) -> dict:
    """Count how often each paragraph fingerprint appears across documents."""
    fingerprint_counts = defaultdict(lambda: {"count": 0, "examples": [], "styles": set()})

    for profile in profiles:
        seen_in_doc = set()  # Deduplicate within same document
        for para in profile["paragraphs"]:
            fp = para["fingerprint"]
            if fp not in seen_in_doc:
                fingerprint_counts[fp]["count"] += 1
                seen_in_doc.add(fp)

            # Keep example text (up to 3)
            if len(fingerprint_counts[fp]["examples"]) < 3:
                if para["text"] not in fingerprint_counts[fp]["examples"]:
                    fingerprint_counts[fp]["examples"].append(para["text"])

            fingerprint_counts[fp]["styles"].add(para["style"])

    # Convert sets to lists for JSON serialization
    for fp in fingerprint_counts:
        fingerprint_counts[fp]["styles"] = list(fingerprint_counts[fp]["styles"])

    return dict(fingerprint_counts)


def classify_paragraphs(fingerprint_counts: dict, total_docs: int, threshold: float) -> dict:
    """
    Phase 2B: Classify fingerprints as fixed (boilerplate) or variable.
    """
    fixed = {}
    variable = {}

    min_count = int(total_docs * threshold)

    for fp, data in fingerprint_counts.items():
        # Check if it's boilerplate
        is_fixed = (
            data["count"] >= min_count and
            len(data["examples"][0]) >= CONFIG["min_text_length"]
        )

        classification = {
            "fingerprint": fp,
            "occurrence_count": data["count"],
            "occurrence_rate": data["count"] / total_docs,
            "examples": data["examples"],
            "styles": data["styles"],
        }

        if is_fixed:
            fixed[fp] = classification
        else:
            variable[fp] = classification

    return {"fixed": fixed, "variable": variable}


def detect_sequences(profiles: list[dict], fixed_fingerprints: set, ngram_sizes: list[int]) -> list[dict]:
    """
    Phase 2C: Detect sequences of paragraphs that appear together.
    """
    sequence_counts = defaultdict(lambda: {"count": 0, "examples": []})

    for profile in profiles:
        fingerprints = [p["fingerprint"] for p in profile["paragraphs"]]

        for n in ngram_sizes:
            for i in range(len(fingerprints) - n + 1):
                seq = tuple(fingerprints[i:i+n])

                # Only count sequences that start with a fixed paragraph
                if seq[0] in fixed_fingerprints:
                    seq_key = "||".join(seq)
                    sequence_counts[seq_key]["count"] += 1

                    if len(sequence_counts[seq_key]["examples"]) < 2:
                        example_texts = [profile["paragraphs"][i+j]["text"][:50] for j in range(n)]
                        sequence_counts[seq_key]["examples"].append(example_texts)

    # Filter to sequences appearing in most documents
    total_docs = len(profiles)
    min_count = int(total_docs * CONFIG["boilerplate_threshold"])

    sequences = []
    for seq_key, data in sequence_counts.items():
        if data["count"] >= min_count:
            sequences.append({
                "sequence": seq_key.split("||"),
                "count": data["count"],
                "rate": data["count"] / total_docs,
                "examples": data["examples"],
            })

    # Sort by length (longer sequences first) then by count
    sequences.sort(key=lambda x: (-len(x["sequence"]), -x["count"]))

    return sequences


# =============================================================================
# PHASE 3: ANCHOR DETECTION & SKELETON BUILDING
# =============================================================================

# Common Gutachten section names (including numbered variants)
KNOWN_ANCHORS = [
    # Main sections (numbered in templates)
    "Gutachterliche Fragestellung",
    "Fragestellung",
    "1. Anamnese",
    "Anamnese",
    "2. Untersuchungsbefunde",
    "Untersuchungsbefunde",
    "3. Diagnosen",
    "Diagnosen",
    "Diagnose",
    "4. Epikrise",
    "Epikrise",
    "5. Sozialmedizinische Leistungsbeurteilung",
    "Sozialmedizinische Leistungsbeurteilung",
    # Sub-sections
    "1.1 Anamnese medizinischer Daten",
    "Anamnese medizinischer Daten",
    "1.2 Biografische Anamnese",
    "Biografische Anamnese",
    "Familienanamnese",
    "Eigenanamnese",
    "Sozialanamnese",
    "Vegetative Anamnese",
    "Aktuelle Beschwerden",
    "Jetzige Beschwerden",
    "Aktuelle Medikation",
    "Therapie und behandelnde Ärzte",
    # Befunde
    "Neurologischer Befund",
    "Körperlicher Befund",
    "Psychischer Befund",
    "Neurologischer/körperlicher Untersuchungsbefund",
    # Other
    "Beurteilung",
    "Zusammenfassung",
    "Therapieempfehlung",
    "Empfehlung",
    "Hamilton Depression Scale",
    "Vorgelegte Klinikaufenthalte",
]


def levenshtein_similarity(s1: str, s2: str) -> float:
    """Compute Levenshtein similarity (1 - normalized distance)."""
    if not s1 or not s2:
        return 0.0

    len1, len2 = len(s1), len(s2)
    if len1 == 0 or len2 == 0:
        return 0.0

    # Create distance matrix
    d = [[0] * (len2 + 1) for _ in range(len1 + 1)]
    for i in range(len1 + 1):
        d[i][0] = i
    for j in range(len2 + 1):
        d[0][j] = j

    for i in range(1, len1 + 1):
        for j in range(1, len2 + 1):
            cost = 0 if s1[i-1] == s2[j-1] else 1
            d[i][j] = min(
                d[i-1][j] + 1,      # deletion
                d[i][j-1] + 1,      # insertion
                d[i-1][j-1] + cost  # substitution
            )

    distance = d[len1][len2]
    max_len = max(len1, len2)
    return 1.0 - (distance / max_len)


def find_anchors(profiles: list[dict], classification: dict) -> list[dict]:
    """
    Phase 3A: Identify anchor paragraphs (section headings).
    """
    # Get all paragraph texts from fixed blocks that look like headings
    heading_candidates = defaultdict(lambda: {"count": 0, "variants": [], "styles": set()})

    for profile in profiles:
        for para in profile["paragraphs"]:
            text = para["text"].strip()
            style = para["style"]

            # Skip empty or very long paragraphs
            if not text or len(text) > 100:
                continue

            # Check if it's a heading style or short line
            is_heading_style = "Heading" in style or style in ["Title", "Subtitle"]
            is_short = len(text) < 60

            if is_heading_style or is_short:
                # Normalize for matching
                norm_text = normalize_basic(text.rstrip(':').rstrip('-').strip())

                # Check against known anchors
                for known in KNOWN_ANCHORS:
                    known_norm = normalize_basic(known)
                    similarity = levenshtein_similarity(norm_text, known_norm)

                    if similarity >= CONFIG["anchor_similarity_threshold"]:
                        heading_candidates[known]["count"] += 1
                        if text not in heading_candidates[known]["variants"]:
                            heading_candidates[known]["variants"].append(text)
                        heading_candidates[known]["styles"].add(style)
                        break

    # Build anchor list
    total_docs = len(profiles)
    anchors = []

    for anchor_name, data in heading_candidates.items():
        if data["count"] >= total_docs * 0.5:  # At least 50% occurrence
            anchors.append({
                "id": anchor_name.lower().replace(" ", "_"),
                "canonical_text": anchor_name,
                "match_mode": "fuzzy",
                "min_similarity": CONFIG["anchor_similarity_threshold"],
                "variants_seen": data["variants"][:5],
                "occurrence_rate": data["count"] / total_docs,
                "styles": list(data["styles"]),
            })

    # Sort by typical Gutachten order
    order = {name.lower().replace(" ", "_"): i for i, name in enumerate(KNOWN_ANCHORS)}
    anchors.sort(key=lambda x: order.get(x["id"], 999))

    return anchors


def build_skeleton(profiles: list[dict], anchors: list[dict], classification: dict) -> list[dict]:
    """
    Phase 3B: Build template skeleton from anchors and fixed blocks.
    """
    skeleton = []
    anchor_ids = {a["id"] for a in anchors}

    # For each anchor, create a slot
    for anchor in anchors:
        # Add fixed block for the anchor heading
        skeleton.append({
            "type": "fixed",
            "id": f"{anchor['id']}_heading",
            "paragraphs": [{
                "text": anchor["canonical_text"],
                "style": anchor["styles"][0] if anchor["styles"] else "Heading 1",
            }],
        })

        # Add slot for content after the anchor
        skeleton.append({
            "type": "slot",
            "slot_id": f"{anchor['id']}_body",
            "section_name": anchor["canonical_text"],
            "allowed_styles": ["BODY", "BULLET"],
            "list_behavior": "bullets_allowed",
            "optional": anchor["occurrence_rate"] < 0.8,
        })

    return skeleton


# =============================================================================
# PHASE 4: STYLE EXTRACTION
# =============================================================================

def extract_style_roles(profiles: list[dict]) -> dict:
    """
    Phase 4A: Determine most common styles for each role.
    """
    # Aggregate style usage
    heading_styles = defaultdict(int)
    body_styles = defaultdict(int)
    list_styles = defaultdict(int)

    for profile in profiles:
        for para in profile["paragraphs"]:
            style = para["style"]
            text = para["text"].strip()

            if "Heading" in style or style in ["Title", "Subtitle"]:
                heading_styles[style] += 1
            elif para["is_list"]:
                list_styles[style] += 1
            elif text and len(text) > 20:
                body_styles[style] += 1

    # Pick most common for each role
    def most_common(d, default):
        if not d:
            return default
        return max(d.items(), key=lambda x: x[1])[0]

    return {
        "H1": most_common({k: v for k, v in heading_styles.items() if "1" in k}, "Heading 1"),
        "H2": most_common({k: v for k, v in heading_styles.items() if "2" in k}, "Heading 2"),
        "H3": most_common({k: v for k, v in heading_styles.items() if "3" in k}, "Heading 3"),
        "BODY": most_common(body_styles, "Normal"),
        "BULLET": most_common(list_styles, "List Bullet"),
        "TITLE": most_common({k: v for k, v in heading_styles.items() if k == "Title"}, "Title"),
    }


def extract_kopfzeile(profiles: list[dict]) -> dict:
    """
    Phase 4C: Extract common header/footer content.
    """
    all_headers = []
    all_footers = []

    for profile in profiles:
        all_headers.extend(profile.get("headers", []))
        all_footers.extend(profile.get("footers", []))

    # Find most common (if any)
    header_counts = defaultdict(int)
    footer_counts = defaultdict(int)

    for h in all_headers:
        header_counts[h] += 1
    for f in all_footers:
        footer_counts[f] += 1

    most_common_header = max(header_counts.items(), key=lambda x: x[1])[0] if header_counts else ""
    most_common_footer = max(footer_counts.items(), key=lambda x: x[1])[0] if footer_counts else ""

    return {
        "header": most_common_header,
        "footer": most_common_footer,
        "header_variants": list(header_counts.keys())[:5],
        "footer_variants": list(footer_counts.keys())[:5],
    }


# =============================================================================
# PHASE 5: OUTPUT TEMPLATE SPEC
# =============================================================================

def build_template_spec(
    profiles: list[dict],
    classification: dict,
    anchors: list[dict],
    skeleton: list[dict],
    style_roles: dict,
    kopfzeile: dict,
    family_id: str = "default",
    family_name: str = "Default Template"
) -> dict:
    """
    Phase 5: Build final TemplateSpec.
    """
    return {
        "version": "1.0",
        "created_at": datetime.now().isoformat(),
        "family_id": family_id,
        "family_name": family_name,

        "anchors": anchors,
        "skeleton": skeleton,
        "style_roles": style_roles,

        "kopfzeile": {
            "content": kopfzeile["header"],
            "variants": kopfzeile["header_variants"],
        },
        "fusszeile": {
            "content": kopfzeile["footer"],
            "variants": kopfzeile["footer_variants"],
        },

        "render_rules": {
            "spacing_after_heading": 12,
            "spacing_after_paragraph": 6,
            "blank_line_before_section": True,
        },

        "quality_metrics": {
            "documents_analyzed": len(profiles),
            "fixed_blocks_found": len(classification["fixed"]),
            "variable_blocks_found": len(classification["variable"]),
            "anchors_detected": len(anchors),
            "boilerplate_threshold": CONFIG["boilerplate_threshold"],
        },
    }


# =============================================================================
# MAIN EXTRACTION PIPELINE
# =============================================================================

def extract_template(input_folder: str, output_folder: str) -> dict:
    """
    Run full extraction pipeline (Phases 0-5).
    """
    input_path = Path(input_folder)
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    # Find all DOCX files
    docx_files = list(input_path.glob("*.docx"))
    if not docx_files:
        print(f"ERROR: No .docx files found in {input_folder}", file=sys.stderr)
        return None

    print(f"Found {len(docx_files)} DOCX files", file=sys.stderr)

    # Phase 0 & 1: Ingest and normalize
    profiles = []
    profiles_dir = output_path / "doc_profiles"
    profiles_dir.mkdir(exist_ok=True)

    for docx_file in docx_files:
        print(f"  Processing: {docx_file.name}", file=sys.stderr)
        try:
            profile = parse_docx_to_profile(str(docx_file))
            profile = add_normalizations(profile)
            profiles.append(profile)

            # Save profile
            profile_path = profiles_dir / f"{docx_file.stem}.json"
            with open(profile_path, 'w', encoding='utf-8') as f:
                json.dump(profile, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"  ERROR processing {docx_file.name}: {e}", file=sys.stderr)

    if not profiles:
        print("ERROR: No profiles extracted", file=sys.stderr)
        return None

    print(f"\nPhase 0-1 complete: {len(profiles)} profiles extracted", file=sys.stderr)

    # Phase 2: Boilerplate detection
    print("Phase 2: Detecting boilerplate...", file=sys.stderr)
    fingerprint_counts = compute_fingerprint_counts(profiles)
    classification = classify_paragraphs(fingerprint_counts, len(profiles), CONFIG["boilerplate_threshold"])

    print(f"  Fixed blocks: {len(classification['fixed'])}", file=sys.stderr)
    print(f"  Variable blocks: {len(classification['variable'])}", file=sys.stderr)

    # Save classification
    with open(output_path / "classification.json", 'w', encoding='utf-8') as f:
        json.dump(classification, f, ensure_ascii=False, indent=2)

    # Sequence detection
    fixed_fps = set(classification["fixed"].keys())
    sequences = detect_sequences(profiles, fixed_fps, CONFIG["ngram_sizes"])
    print(f"  Sequences found: {len(sequences)}", file=sys.stderr)

    with open(output_path / "sequences.json", 'w', encoding='utf-8') as f:
        json.dump(sequences, f, ensure_ascii=False, indent=2)

    # Phase 3: Anchors and skeleton
    print("Phase 3: Building skeleton...", file=sys.stderr)
    anchors = find_anchors(profiles, classification)
    print(f"  Anchors found: {[a['canonical_text'] for a in anchors]}", file=sys.stderr)

    skeleton = build_skeleton(profiles, anchors, classification)

    # Phase 4: Style extraction
    print("Phase 4: Extracting styles...", file=sys.stderr)
    style_roles = extract_style_roles(profiles)
    kopfzeile = extract_kopfzeile(profiles)
    print(f"  Style roles: {style_roles}", file=sys.stderr)

    # Phase 5: Build template spec
    print("Phase 5: Building template spec...", file=sys.stderr)
    template_spec = build_template_spec(
        profiles, classification, anchors, skeleton,
        style_roles, kopfzeile
    )

    # Save template spec
    spec_path = output_path / "template_spec.json"
    with open(spec_path, 'w', encoding='utf-8') as f:
        json.dump(template_spec, f, ensure_ascii=False, indent=2)

    print(f"\nTemplate extraction complete!", file=sys.stderr)
    print(f"Output: {spec_path}", file=sys.stderr)

    return template_spec


# =============================================================================
# CLI
# =============================================================================

def main():
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python template_extractor.py extract <folder_with_docx> [output_folder]")
        print("  python template_extractor.py analyze <doc_profiles_folder>")
        sys.exit(1)

    command = sys.argv[1]
    input_folder = sys.argv[2]

    if command == "extract":
        output_folder = sys.argv[3] if len(sys.argv) > 3 else "template_output"
        result = extract_template(input_folder, output_folder)
        if result:
            print(json.dumps(result, ensure_ascii=False, indent=2))

    elif command == "analyze":
        # Load existing profiles and re-run analysis
        print("Analyze mode not yet implemented", file=sys.stderr)

    else:
        print(f"Unknown command: {command}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
