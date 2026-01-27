"""
DOCX Renderer - Deterministic document generation

Takes template_spec.json + content.json → Final Gutachten.docx

Features:
- Walks template skeleton in order
- Writes FixedBlocks exactly as specified
- Fills Slots with LLM-structured content
- Applies yellow highlight to {unclear:...} spans
- Handles missing slots gracefully
- Inserts Kopfzeile/Fußzeile
"""

import sys
import os
import json
import re
from datetime import datetime
from pathlib import Path

# Force UTF-8 for Windows
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# =============================================================================
# HIGHLIGHT COLORS
# =============================================================================

# Yellow highlight for unclear text (WD_COLOR_INDEX.YELLOW = 7)
HIGHLIGHT_YELLOW = 7


def add_highlight(run, color_index=HIGHLIGHT_YELLOW):
    """Add highlight color to a run."""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


# =============================================================================
# DOCX RENDERER CLASS
# =============================================================================

class DocxRenderer:
    def __init__(self, template_spec: dict, base_template_path: str = None):
        """
        Initialize renderer with template spec.

        Args:
            template_spec: Loaded template_spec.json
            base_template_path: Path to base_template.docx (optional)
        """
        self.template_spec = template_spec
        self.base_template_path = base_template_path

        # Extract style mappings
        self.style_roles = template_spec.get("style_roles", {
            "H1": "Heading 1",
            "H2": "Heading 2",
            "BODY": "Normal",
            "BULLET": "List Bullet",
        })

        # Extract render rules
        self.render_rules = template_spec.get("render_rules", {
            "spacing_after_heading": 12,
            "spacing_after_paragraph": 6,
            "blank_line_before_section": True,
        })

    def render(self, content: dict, output_path: str) -> str:
        """
        Render final DOCX from content.json.

        Args:
            content: Loaded content.json with slots, unclear_spans, missing_slots
            output_path: Where to save the final .docx

        Returns:
            Path to generated document
        """
        # Create document (from base template or blank)
        if self.base_template_path and os.path.exists(self.base_template_path):
            doc = Document(self.base_template_path)
        else:
            doc = Document()
            self._setup_default_styles(doc)

        # Add Kopfzeile (header) if specified
        self._add_header(doc)

        # Walk skeleton and render
        skeleton = self.template_spec.get("skeleton", [])
        slots = content.get("slots", {})
        missing_slots = set(content.get("missing_slots", []))

        for item in skeleton:
            item_type = item.get("type")

            if item_type == "fixed":
                self._render_fixed_block(doc, item)

            elif item_type == "slot":
                slot_id = item.get("slot_id")

                if slot_id in missing_slots:
                    # Optionally add placeholder for missing section
                    if not item.get("optional", False):
                        self._render_missing_slot(doc, item)
                elif slot_id in slots:
                    self._render_slot(doc, item, slots[slot_id], content.get("unclear_spans", []))

        # Add Fußzeile (footer) if specified
        self._add_footer(doc)

        # Save document
        doc.save(output_path)
        print(f"[RENDERER] Saved: {output_path}", file=sys.stderr)

        return output_path

    def _setup_default_styles(self, doc):
        """Set up default styles for blank document."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def _add_header(self, doc):
        """Add Kopfzeile to document."""
        kopfzeile = self.template_spec.get("kopfzeile", {})
        header_content = kopfzeile.get("content", "")

        if header_content:
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_content

    def _add_footer(self, doc):
        """Add Fußzeile to document."""
        fusszeile = self.template_spec.get("fusszeile", {})
        footer_content = fusszeile.get("content", "")

        if footer_content:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = footer_content

    def _render_fixed_block(self, doc, item):
        """Render a fixed (boilerplate) block."""
        paragraphs = item.get("paragraphs", [])

        for para_spec in paragraphs:
            text = para_spec.get("text", "")
            style_role = para_spec.get("style", "BODY")
            style_name = self.style_roles.get(style_role, style_role)

            para = doc.add_paragraph(text)
            try:
                para.style = style_name
            except KeyError:
                # Style not found, use Normal
                para.style = 'Normal'

    def _render_slot(self, doc, slot_spec, paragraphs: list, unclear_spans: list):
        """Render a slot with content from LLM."""
        slot_id = slot_spec.get("slot_id")
        allowed_styles = slot_spec.get("allowed_styles", ["BODY"])
        list_behavior = slot_spec.get("list_behavior", "none")

        # Find unclear spans for this slot
        slot_unclear = [u for u in unclear_spans if u.get("slot_id") == slot_id]
        unclear_texts = {u.get("text") for u in slot_unclear}

        for para_text in paragraphs:
            if not para_text or not para_text.strip():
                continue

            # Check if this is a bullet point
            is_bullet = (
                list_behavior in ["bullets_allowed", "bullets_only"] and
                (para_text.strip().startswith("- ") or
                 para_text.strip().startswith("• ") or
                 para_text.strip().startswith("* "))
            )

            # Determine style
            if is_bullet:
                style_name = self.style_roles.get("BULLET", "List Bullet")
                para_text = para_text.strip().lstrip("-•* ").strip()
            else:
                style_name = self.style_roles.get(allowed_styles[0], "Normal")

            # Create paragraph
            para = doc.add_paragraph()
            try:
                para.style = style_name
            except KeyError:
                para.style = 'Normal'

            # Handle {unclear:...} highlighting
            self._add_text_with_highlights(para, para_text, unclear_texts)

    def _add_text_with_highlights(self, para, text: str, unclear_texts: set):
        """Add text to paragraph, highlighting unclear parts."""
        # Find all {unclear:...} patterns
        pattern = r'\{unclear:([^}]+)\}'
        last_end = 0

        for match in re.finditer(pattern, text):
            # Add text before match
            if match.start() > last_end:
                para.add_run(text[last_end:match.start()])

            # Add highlighted text
            unclear_text = match.group(1)
            run = para.add_run(unclear_text)
            add_highlight(run)

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            remaining = text[last_end:]

            # Also check for unclear texts that weren't marked with {unclear:}
            for unclear in unclear_texts:
                if unclear in remaining and f"{{unclear:{unclear}}}" not in text:
                    # Highlight this unclear text
                    parts = remaining.split(unclear)
                    for i, part in enumerate(parts):
                        if part:
                            para.add_run(part)
                        if i < len(parts) - 1:
                            run = para.add_run(unclear)
                            add_highlight(run)
                    return

            para.add_run(remaining)

    def _render_missing_slot(self, doc, slot_spec):
        """Render placeholder for missing (non-optional) slot."""
        section_name = slot_spec.get("section_name", slot_spec.get("slot_id", ""))

        para = doc.add_paragraph()
        run = para.add_run(f"[Abschnitt '{section_name}' fehlt im Diktat]")
        run.italic = True
        add_highlight(run)


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def render_gutachten(
    template_spec_path: str,
    content_path: str,
    output_path: str,
    base_template_path: str = None
) -> str:
    """
    Convenience function to render a Gutachten.

    Args:
        template_spec_path: Path to template_spec.json
        content_path: Path to content.json
        output_path: Path for output .docx
        base_template_path: Optional path to base_template.docx

    Returns:
        Path to generated document
    """
    # Load template spec
    with open(template_spec_path, 'r', encoding='utf-8') as f:
        template_spec = json.load(f)

    # Load content
    with open(content_path, 'r', encoding='utf-8') as f:
        content = json.load(f)

    # Render
    renderer = DocxRenderer(template_spec, base_template_path)
    return renderer.render(content, output_path)


def create_base_template(template_spec: dict, output_path: str):
    """
    Create a base_template.docx from template_spec.

    This creates an empty document with all required styles defined.
    """
    doc = Document()

    # Set up page margins (A4 standard)
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set up default Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Set up Heading styles
    style_roles = template_spec.get("style_roles", {})

    for role, style_name in style_roles.items():
        if "Heading" in style_name:
            try:
                heading_style = doc.styles[style_name]
                heading_style.font.name = 'Arial'
                heading_style.font.bold = True

                if "1" in style_name:
                    heading_style.font.size = Pt(14)
                elif "2" in style_name:
                    heading_style.font.size = Pt(12)
                else:
                    heading_style.font.size = Pt(11)
            except KeyError:
                pass

    # Add header/footer
    kopfzeile = template_spec.get("kopfzeile", {})
    if kopfzeile.get("content"):
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = kopfzeile["content"]

    fusszeile = template_spec.get("fusszeile", {})
    if fusszeile.get("content"):
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = fusszeile["content"]

    # Save
    doc.save(output_path)
    print(f"[RENDERER] Created base template: {output_path}", file=sys.stderr)


# =============================================================================
# CLI
# =============================================================================

def main():
    if len(sys.argv) < 4:
        print("Usage:")
        print("  python docx_renderer.py render <template_spec.json> <content.json> <output.docx>")
        print("  python docx_renderer.py create-base <template_spec.json> <base_template.docx>")
        print("  python docx_renderer.py --test")
        sys.exit(1)

    command = sys.argv[1]

    if command == "render":
        template_spec_path = sys.argv[2]
        content_path = sys.argv[3]
        output_path = sys.argv[4]
        base_template = sys.argv[5] if len(sys.argv) > 5 else None

        result = render_gutachten(template_spec_path, content_path, output_path, base_template)
        print(f"Generated: {result}")

    elif command == "create-base":
        template_spec_path = sys.argv[2]
        output_path = sys.argv[3]

        with open(template_spec_path, 'r', encoding='utf-8') as f:
            template_spec = json.load(f)

        create_base_template(template_spec, output_path)

    elif command == "--test":
        # Create a test template spec and content
        test_template = {
            "skeleton": [
                {"type": "fixed", "paragraphs": [{"text": "GUTACHTEN", "style": "H1"}]},
                {"type": "slot", "slot_id": "fragestellung_body", "section_name": "Fragestellung", "allowed_styles": ["BODY"], "optional": False},
                {"type": "fixed", "paragraphs": [{"text": "1. Anamnese", "style": "H2"}]},
                {"type": "slot", "slot_id": "anamnese_body", "section_name": "Anamnese", "allowed_styles": ["BODY", "BULLET"], "list_behavior": "bullets_allowed", "optional": False},
            ],
            "style_roles": {"H1": "Heading 1", "H2": "Heading 2", "BODY": "Normal", "BULLET": "List Bullet"},
            "kopfzeile": {"content": "Dr. med. Test\nMusterstraße 1"},
            "fusszeile": {"content": ""},
            "render_rules": {},
        }

        test_content = {
            "slots": {
                "fragestellung_body": ["Die Begutachtung erfolgt auf Veranlassung der DRV."],
                "anamnese_body": [
                    "Der Patient berichtet über {unclear:Beschwerden im linken Bereich} seit 2019.",
                    "- Rückenschmerzen",
                    "- Kopfschmerzen",
                ],
            },
            "unclear_spans": [
                {"slot_id": "anamnese_body", "text": "Beschwerden im linken Bereich", "reason": "garbled"}
            ],
            "missing_slots": [],
        }

        renderer = DocxRenderer(test_template)
        renderer.render(test_content, "test_output.docx")
        print("Test document saved to: test_output.docx")

    else:
        print(f"Unknown command: {command}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
