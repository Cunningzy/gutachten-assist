"""
DOCX Formatter Module
Applies FormatSpec to DOCX documents using python-docx.
"""

import os
import copy
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from format_spec import FormatSpec, FontSpec, ParagraphSpec, StyleSpec, validate_spec


@dataclass
class FormattingResult:
    """Result of a formatting operation"""
    success: bool
    input_file: str
    output_file: str
    applied_changes: Dict[str, int] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'success': self.success,
            'input_file': self.input_file,
            'output_file': self.output_file,
            'applied_changes': self.applied_changes,
            'warnings': self.warnings,
            'errors': self.errors
        }


class DocxFormatter:
    """
    Applies formatting specifications to DOCX documents.
    """

    def __init__(self, spec: FormatSpec):
        self.spec = spec
        self.warnings: List[str] = []
        self.changes: Dict[str, int] = {
            'paragraphs_styled': 0,
            'runs_formatted': 0,
            'styles_updated': 0,
            'tables_formatted': 0,
            'headers_updated': 0,
            'footers_updated': 0,
            'margins_set': 0
        }

    def format_document(self, input_path: str, output_path: str) -> FormattingResult:
        """
        Apply formatting spec to a document.

        Args:
            input_path: Path to input DOCX file
            output_path: Path to save formatted DOCX file

        Returns:
            FormattingResult with details of changes made
        """
        result = FormattingResult(
            success=False,
            input_file=input_path,
            output_file=output_path
        )

        try:
            # Load document
            if not os.path.exists(input_path):
                result.errors.append(f"Input file not found: {input_path}")
                return result

            doc = Document(input_path)

            # Apply formatting based on scope
            scope = self.spec.scope

            if scope in ['whole_document', 'styles_only']:
                self._apply_styles(doc)

            if scope in ['whole_document', 'body_only']:
                self._apply_defaults_to_body(doc)
                self._apply_page_settings(doc)
                self._apply_headers_footers(doc)

            if scope in ['whole_document', 'tables_only']:
                self._apply_table_formatting(doc)

            # Apply title page if specified
            if self.spec.title_page and self.spec.title_page.enabled:
                self._create_title_page(doc)

            # Save document
            doc.save(output_path)

            result.success = True
            result.applied_changes = self.changes.copy()
            result.warnings = self.warnings.copy()

        except Exception as e:
            result.errors.append(f"Error formatting document: {str(e)}")

        return result

    def _apply_styles(self, doc: Document):
        """Apply style definitions from spec"""
        if not self.spec.styles:
            return

        for style_name, style_spec in self.spec.styles.items():
            if not style_spec:
                continue

            try:
                # Try to get existing style
                style = None
                for s in doc.styles:
                    if s.name == style_name:
                        style = s
                        break

                if style is None:
                    self.warnings.append(f"Style '{style_name}' not found in document, skipping")
                    continue

                # Apply font settings
                if style_spec.font:
                    self._apply_font_to_style(style, style_spec.font)

                # Apply paragraph settings (for paragraph styles)
                if style_spec.paragraph and hasattr(style, 'paragraph_format'):
                    self._apply_paragraph_format(style.paragraph_format, style_spec.paragraph)

                self.changes['styles_updated'] += 1

            except Exception as e:
                self.warnings.append(f"Could not update style '{style_name}': {str(e)}")

    def _apply_font_to_style(self, style, font_spec: FontSpec):
        """Apply font settings to a style"""
        if not hasattr(style, 'font'):
            return

        font = style.font

        if font_spec.name:
            font.name = font_spec.name
        if font_spec.size_pt:
            font.size = Pt(font_spec.size_pt)
        if font_spec.bold is not None:
            font.bold = font_spec.bold
        if font_spec.italic is not None:
            font.italic = font_spec.italic
        if font_spec.underline is not None:
            font.underline = font_spec.underline
        if font_spec.color_hex:
            try:
                r = int(font_spec.color_hex[0:2], 16)
                g = int(font_spec.color_hex[2:4], 16)
                b = int(font_spec.color_hex[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)
            except (ValueError, IndexError):
                self.warnings.append(f"Invalid color hex: {font_spec.color_hex}")

    def _apply_paragraph_format(self, para_format, para_spec: ParagraphSpec):
        """Apply paragraph formatting"""
        if para_spec.line_spacing:
            para_format.line_spacing = para_spec.line_spacing

        if para_spec.space_before_pt is not None:
            para_format.space_before = Pt(para_spec.space_before_pt)

        if para_spec.space_after_pt is not None:
            para_format.space_after = Pt(para_spec.space_after_pt)

        if para_spec.alignment:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if para_spec.alignment in alignment_map:
                para_format.alignment = alignment_map[para_spec.alignment]

        if para_spec.first_line_indent_mm is not None:
            para_format.first_line_indent = Mm(para_spec.first_line_indent_mm)

    def _apply_defaults_to_body(self, doc: Document):
        """Apply default formatting to body paragraphs"""
        if not self.spec.defaults:
            return

        for paragraph in doc.paragraphs:
            # Apply paragraph formatting
            if self.spec.defaults.paragraph:
                self._apply_paragraph_format(paragraph.paragraph_format, self.spec.defaults.paragraph)

            # Apply font formatting to runs
            if self.spec.defaults.font:
                for run in paragraph.runs:
                    self._apply_font_to_run(run, self.spec.defaults.font)
                    self.changes['runs_formatted'] += 1

            self.changes['paragraphs_styled'] += 1

    def _apply_font_to_run(self, run, font_spec: FontSpec, force: bool = False):
        """Apply font settings to a run"""
        font = run.font

        # Only apply if not already set (unless force=True)
        if font_spec.name and (force or font.name is None):
            font.name = font_spec.name
            # Also set for complex scripts (German umlauts etc.)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_spec.name)

        if font_spec.size_pt and (force or font.size is None):
            font.size = Pt(font_spec.size_pt)

        if font_spec.bold is not None and (force or font.bold is None):
            font.bold = font_spec.bold

        if font_spec.italic is not None and (force or font.italic is None):
            font.italic = font_spec.italic

        if font_spec.underline is not None and (force or font.underline is None):
            font.underline = font_spec.underline

        if font_spec.color_hex:
            try:
                r = int(font_spec.color_hex[0:2], 16)
                g = int(font_spec.color_hex[2:4], 16)
                b = int(font_spec.color_hex[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass

    def _apply_page_settings(self, doc: Document):
        """Apply page layout settings"""
        if not self.spec.page:
            return

        for section in doc.sections:
            # Apply margins
            if self.spec.page.margins:
                margins = self.spec.page.margins
                if margins.top_mm is not None:
                    section.top_margin = Mm(margins.top_mm)
                if margins.bottom_mm is not None:
                    section.bottom_margin = Mm(margins.bottom_mm)
                if margins.left_mm is not None:
                    section.left_margin = Mm(margins.left_mm)
                if margins.right_mm is not None:
                    section.right_margin = Mm(margins.right_mm)
                self.changes['margins_set'] += 1

            # Apply orientation
            if self.spec.page.orientation:
                if self.spec.page.orientation == 'landscape':
                    section.orientation = WD_ORIENT.LANDSCAPE
                    # Swap dimensions for landscape
                    new_width = section.page_height
                    new_height = section.page_width
                    section.page_width = new_width
                    section.page_height = new_height
                elif self.spec.page.orientation == 'portrait':
                    section.orientation = WD_ORIENT.PORTRAIT

            # Apply page size
            if self.spec.page.size:
                if self.spec.page.size == 'A4':
                    section.page_width = Mm(210)
                    section.page_height = Mm(297)
                elif self.spec.page.size == 'Letter':
                    section.page_width = Inches(8.5)
                    section.page_height = Inches(11)

    def _apply_headers_footers(self, doc: Document):
        """Apply header and footer settings"""
        for section in doc.sections:
            # Apply header
            if self.spec.header and self.spec.header.enabled:
                self._apply_header(section, self.spec.header)
                self.changes['headers_updated'] += 1

            # Apply footer
            if self.spec.footer and self.spec.footer.enabled:
                self._apply_footer(section, self.spec.footer)
                self.changes['footers_updated'] += 1

    def _apply_header(self, section, header_spec):
        """Apply header to a section"""
        header = section.header
        header.is_linked_to_previous = False

        if header_spec.content:
            # Clear existing content
            for para in header.paragraphs:
                para.clear()

            content = header_spec.content

            # Create a table for left/center/right alignment
            if content.left_text or content.center_text or content.right_text:
                # Use first paragraph
                para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                para.clear()

                # Simple approach: use tabs for positioning
                if content.center_text:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(content.center_text)
                    if content.font:
                        self._apply_font_to_run(run, content.font, force=True)
                elif content.left_text:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = para.add_run(content.left_text)
                    if content.font:
                        self._apply_font_to_run(run, content.font, force=True)
                elif content.right_text:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = para.add_run(content.right_text)
                    if content.font:
                        self._apply_font_to_run(run, content.font, force=True)

    def _apply_footer(self, section, footer_spec):
        """Apply footer to a section"""
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing content
        for para in footer.paragraphs:
            para.clear()

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()

        # Add page number if enabled
        if footer_spec.page_number and footer_spec.page_number.enabled:
            position = footer_spec.page_number.position
            if position == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif position == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add page number field
            pn_format = footer_spec.page_number.format
            if pn_format in ['1', '1/3', 'Page 1 of 3', 'Seite 1 von 3']:
                self._add_page_number_field(para, pn_format)

        elif footer_spec.content:
            content = footer_spec.content
            if content.center_text:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(content.center_text)
                if content.font:
                    self._apply_font_to_run(run, content.font, force=True)
            elif content.left_text:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(content.left_text)
                if content.font:
                    self._apply_font_to_run(run, content.font, force=True)
            elif content.right_text:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = para.add_run(content.right_text)
                if content.font:
                    self._apply_font_to_run(run, content.font, force=True)

    def _add_page_number_field(self, paragraph, format_str: str):
        """Add page number field to paragraph"""
        run = paragraph.add_run()

        # Create PAGE field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = " PAGE "

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

        # Add " of X" or " von X" if needed
        if format_str in ['1/3', 'Page 1 of 3', 'Seite 1 von 3']:
            if format_str == '1/3':
                paragraph.add_run('/')
            elif format_str == 'Page 1 of 3':
                paragraph.add_run(' of ')
            else:
                paragraph.add_run(' von ')

            # Add NUMPAGES field
            run2 = paragraph.add_run()
            fld_char_begin2 = OxmlElement('w:fldChar')
            fld_char_begin2.set(qn('w:fldCharType'), 'begin')

            instr_text2 = OxmlElement('w:instrText')
            instr_text2.set(qn('xml:space'), 'preserve')
            instr_text2.text = " NUMPAGES "

            fld_char_end2 = OxmlElement('w:fldChar')
            fld_char_end2.set(qn('w:fldCharType'), 'end')

            run2._r.append(fld_char_begin2)
            run2._r.append(instr_text2)
            run2._r.append(fld_char_end2)

    def _apply_table_formatting(self, doc: Document):
        """Apply formatting to tables"""
        if not self.spec.tables:
            return

        tables_spec = self.spec.tables

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            # Apply default font
                            if tables_spec.default_font_name:
                                run.font.name = tables_spec.default_font_name
                            if tables_spec.default_font_size_pt:
                                run.font.size = Pt(tables_spec.default_font_size_pt)

                            # Apply bold to header row
                            if row_idx == 0 and tables_spec.header_row_bold:
                                run.font.bold = True

            self.changes['tables_formatted'] += 1

    def _create_title_page(self, doc: Document):
        """Create a title page at the beginning of the document"""
        if not self.spec.title_page:
            return

        tp = self.spec.title_page

        # Insert paragraphs at the beginning
        # We'll insert in reverse order since we're prepending

        # Add page break after title content
        if tp.add_page_break_after:
            first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
            run = first_para.insert_paragraph_before().add_run()
            run.add_break()  # Page break

        # Add date
        if tp.date_text:
            para = doc.paragraphs[0].insert_paragraph_before()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(tp.date_text)
            run.font.size = Pt(12)

        # Add author
        if tp.author_text:
            para = doc.paragraphs[0].insert_paragraph_before()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(tp.author_text)
            run.font.size = Pt(12)

        # Add subtitle
        if tp.subtitle_text:
            para = doc.paragraphs[0].insert_paragraph_before()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(tp.subtitle_text)
            run.font.size = Pt(14)

        # Add title (with optional custom font)
        if tp.title_text:
            para = doc.paragraphs[0].insert_paragraph_before()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(tp.title_text)

            if tp.title_font:
                self._apply_font_to_run(run, tp.title_font, force=True)
            else:
                run.font.size = Pt(18)
                run.font.bold = True

        # Add some spacing at top
        para = doc.paragraphs[0].insert_paragraph_before()
        para.add_run()  # Empty paragraph for spacing


def format_docx(input_path: str, output_path: str, spec: FormatSpec) -> FormattingResult:
    """
    Convenience function to format a DOCX file.

    Args:
        input_path: Path to input DOCX
        output_path: Path to save formatted DOCX
        spec: FormatSpec to apply

    Returns:
        FormattingResult
    """
    formatter = DocxFormatter(spec)
    return formatter.format_document(input_path, output_path)


def format_docx_from_json(input_path: str, output_path: str, spec_json: str) -> FormattingResult:
    """
    Format a DOCX file using a JSON spec string.

    Args:
        input_path: Path to input DOCX
        output_path: Path to save formatted DOCX
        spec_json: JSON string containing FormatSpec

    Returns:
        FormattingResult
    """
    spec = FormatSpec.from_json(spec_json)
    return format_docx(input_path, output_path, spec)
