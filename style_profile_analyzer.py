"""
Style Profile Analyzer for Gutachten Assistant
Analyzes multiple example Gutachten to extract a reusable TEMPLATE.

This follows the proper approach:
1. Extract paragraphs with full formatting metadata from each DOCX
2. Create fingerprints for matching (STRICT or COMMON mode)
3. Compute intersection - what appears in ALL documents
4. Build a template DOCX preserving formatting from reference document

The LLM is NOT used to generate DOCX - only deterministic code does that.

Usage:
    python style_profile_analyzer.py <documents_json> <output_path>
"""

import sys
import os
import json
import re
import hashlib
from datetime import datetime
from collections import Counter, defaultdict
from typing import List, Dict, Set, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy


class ParagraphFingerprint:
    """Represents a paragraph with its text and formatting for matching."""

    def __init__(self, para, block_id: str, mode: str = "COMMON"):
        self.block_id = block_id
        self.raw_text = para.text
        self.norm_text = self._normalize_text(para.text)
        self.style_name = para.style.name if para.style else "Normal"
        self.mode = mode

        # Extract formatting
        self.alignment = str(para.alignment) if para.alignment else "LEFT"
        self.runs_info = []
        for run in para.runs:
            self.runs_info.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': str(run.font.size) if run.font.size else None
            })

        # Compute fingerprint hash
        if mode == "STRICT":
            self.fingerprint = self._hash(self.raw_text)
        else:  # COMMON mode
            self.fingerprint = self._hash(self.norm_text)

    def _normalize_text(self, text: str) -> str:
        """Normalize text for COMMON mode matching."""
        if not text:
            return ""
        # Collapse whitespace
        text = re.sub(r'\s+', ' ', text)
        # Normalize non-breaking spaces
        text = text.replace('\u00a0', ' ')
        # Trim
        text = text.strip()
        # Normalize trailing colons in headings
        text = re.sub(r':$', '', text)
        return text

    def _hash(self, text: str) -> str:
        """Create a stable hash for matching."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()

    def is_likely_patient_specific(self) -> bool:
        """Heuristic to detect patient-specific content."""
        text = self.norm_text.lower()

        # Patterns that suggest patient-specific content
        patient_patterns = [
            r'\b(herr|frau|patient|patientin)\s+[A-ZÄÖÜ][a-zäöüß]+',  # Names
            r'\bgeb\.\s*\d',  # Birth dates
            r'\bgeboren\s+(am\s*)?\d',
            r'\bversicherungsnummer\s*:?\s*\d',
            r'\baz\.?\s*:?\s*\d',  # Aktenzeichen
            r'\bwohnhaft',
            r'\bstraße|str\.',
            r'\d{5}\s+[A-ZÄÖÜ]',  # PLZ + City
            r'\b\d{1,2}\.\d{1,2}\.\d{2,4}\b',  # Dates
        ]

        for pattern in patient_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True

        return False

    def to_dict(self) -> dict:
        return {
            'block_id': self.block_id,
            'raw_text': self.raw_text,
            'norm_text': self.norm_text,
            'style': self.style_name,
            'fingerprint': self.fingerprint,
            'alignment': self.alignment,
            'runs': self.runs_info,
            'is_patient_specific': self.is_likely_patient_specific()
        }


class DocumentExtractor:
    """Extracts structured data from a DOCX file."""

    def __init__(self, doc_path: str, doc_id: str):
        self.doc_path = doc_path
        self.doc_id = doc_id
        self.doc = Document(doc_path)
        self.paragraphs: List[ParagraphFingerprint] = []
        self.header_paragraphs: List[ParagraphFingerprint] = []
        self.footer_paragraphs: List[ParagraphFingerprint] = []

    def extract(self, mode: str = "COMMON") -> dict:
        """Extract all content with formatting metadata."""

        # Extract body paragraphs
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip():  # Skip empty paragraphs
                fp = ParagraphFingerprint(para, f"{self.doc_id}_body_{i}", mode)
                self.paragraphs.append(fp)

        # Extract headers
        for section in self.doc.sections:
            header = section.header
            if header:
                for i, para in enumerate(header.paragraphs):
                    if para.text.strip():
                        fp = ParagraphFingerprint(para, f"{self.doc_id}_header_{i}", mode)
                        self.header_paragraphs.append(fp)

            # Extract footers
            footer = section.footer
            if footer:
                for i, para in enumerate(footer.paragraphs):
                    if para.text.strip():
                        fp = ParagraphFingerprint(para, f"{self.doc_id}_footer_{i}", mode)
                        self.footer_paragraphs.append(fp)

        return {
            'doc_id': self.doc_id,
            'doc_path': self.doc_path,
            'body': [p.to_dict() for p in self.paragraphs],
            'header': [p.to_dict() for p in self.header_paragraphs],
            'footer': [p.to_dict() for p in self.footer_paragraphs],
            'paragraph_count': len(self.paragraphs),
            'fingerprints': {p.fingerprint for p in self.paragraphs}
        }


class TemplateBuilder:
    """Builds a template DOCX from the intersection of multiple documents."""

    def __init__(self, reference_doc_path: str):
        """Use reference document as the formatting source."""
        self.reference_doc = Document(reference_doc_path)
        self.reference_path = reference_doc_path

    def build_template(self, keep_fingerprints: Set[str], output_path: str,
                       clear_header: bool = True, clear_footer: bool = True,
                       mode: str = "COMMON") -> str:
        """
        Build a new template document containing only the common paragraphs.

        Args:
            keep_fingerprints: Set of fingerprints to keep
            output_path: Where to save the template
            clear_header: Whether to clear the header (recommended for patient-specific headers)
            clear_footer: Whether to clear the footer
            mode: STRICT or COMMON matching mode
        """
        # Create new document
        new_doc = Document()

        # Copy section properties (margins, page size) from reference
        ref_section = self.reference_doc.sections[0]
        new_section = new_doc.sections[0]

        new_section.page_width = ref_section.page_width
        new_section.page_height = ref_section.page_height
        new_section.left_margin = ref_section.left_margin
        new_section.right_margin = ref_section.right_margin
        new_section.top_margin = ref_section.top_margin
        new_section.bottom_margin = ref_section.bottom_margin

        # Process header
        if not clear_header:
            # Copy header from reference
            ref_header = ref_section.header
            new_header = new_section.header
            for para in ref_header.paragraphs:
                fp = ParagraphFingerprint(para, "header", mode)
                if fp.fingerprint in keep_fingerprints or not fp.is_likely_patient_specific():
                    self._copy_paragraph(para, new_header.paragraphs[-1] if new_header.paragraphs else new_header.add_paragraph())

        # Process body paragraphs
        kept_count = 0
        for para in self.reference_doc.paragraphs:
            if not para.text.strip():
                # Keep empty paragraphs for spacing
                new_doc.add_paragraph()
                continue

            fp = ParagraphFingerprint(para, "body", mode)

            if fp.fingerprint in keep_fingerprints and not fp.is_likely_patient_specific():
                # This paragraph is common to all docs - keep it
                new_para = new_doc.add_paragraph()
                self._copy_paragraph(para, new_para)
                kept_count += 1
            else:
                # This is document-specific content - add placeholder or skip
                # For now, we add a placeholder marker
                if fp.is_likely_patient_specific():
                    # Add placeholder for patient-specific content
                    placeholder = self._generate_placeholder(fp)
                    if placeholder:
                        new_para = new_doc.add_paragraph(placeholder)
                        new_para.style = para.style

        # Save the template
        new_doc.save(output_path)
        print(f"Template saved: {output_path} ({kept_count} common paragraphs)", file=sys.stderr)

        return output_path

    def _copy_paragraph(self, source_para, target_para):
        """Copy a paragraph with its formatting."""
        # Copy style
        if source_para.style:
            try:
                target_para.style = source_para.style
            except:
                pass

        # Copy alignment
        target_para.alignment = source_para.alignment

        # Copy paragraph formatting
        if source_para.paragraph_format:
            pf = source_para.paragraph_format
            tf = target_para.paragraph_format
            if pf.space_before:
                tf.space_before = pf.space_before
            if pf.space_after:
                tf.space_after = pf.space_after
            if pf.line_spacing:
                tf.line_spacing = pf.line_spacing
            if pf.first_line_indent:
                tf.first_line_indent = pf.first_line_indent
            if pf.left_indent:
                tf.left_indent = pf.left_indent

        # Copy runs with formatting
        for run in source_para.runs:
            new_run = target_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size

    def _generate_placeholder(self, fp: ParagraphFingerprint) -> Optional[str]:
        """Generate a placeholder for patient-specific content."""
        text = fp.norm_text.lower()

        if 'name' in text or 'herr' in text or 'frau' in text:
            return "<<PATIENT_NAME>>"
        if 'geb' in text or 'geboren' in text:
            return "<<GEBURTSDATUM>>"
        if 'wohnhaft' in text or 'adresse' in text:
            return "<<ADRESSE>>"
        if 'versicherungsnummer' in text:
            return "<<VERSICHERUNGSNUMMER>>"
        if 'az' in text or 'aktenzeichen' in text:
            return "<<AKTENZEICHEN>>"

        return None  # Skip this paragraph entirely


class StyleProfileAnalyzer:
    """Main analyzer that coordinates the template extraction process."""

    def __init__(self, mode: str = "COMMON"):
        self.mode = mode
        self.documents: List[dict] = []

    def analyze_documents(self, doc_paths: List[str]) -> dict:
        """Analyze documents and compute the intersection."""
        print(f"Analyzing {len(doc_paths)} documents in {self.mode} mode...", file=sys.stderr)

        # Step 1: Extract all documents
        for i, doc_path in enumerate(doc_paths):
            print(f"  Extracting {i+1}/{len(doc_paths)}: {os.path.basename(doc_path)}", file=sys.stderr)
            try:
                extractor = DocumentExtractor(doc_path, f"doc{i+1}")
                doc_data = extractor.extract(self.mode)
                self.documents.append(doc_data)
            except Exception as e:
                print(f"    Error: {e}", file=sys.stderr)

        if len(self.documents) < 2:
            raise ValueError("Need at least 2 documents to find common content")

        # Step 2: Compute intersection of fingerprints
        all_fingerprint_sets = [doc['fingerprints'] for doc in self.documents]
        common_fingerprints = set.intersection(*all_fingerprint_sets)

        print(f"Found {len(common_fingerprints)} common paragraphs across all documents", file=sys.stderr)

        # Step 3: Analyze what's common
        # Use first document as reference
        reference_doc = self.documents[0]

        # Build list of common blocks
        common_blocks = []
        excluded_blocks = []

        for block in reference_doc['body']:
            if block['fingerprint'] in common_fingerprints:
                if block['is_patient_specific']:
                    excluded_blocks.append({
                        'block_id': block['block_id'],
                        'text_preview': block['raw_text'][:50],
                        'reason': 'patient_specific_content'
                    })
                else:
                    common_blocks.append({
                        'block_id': block['block_id'],
                        'text': block['raw_text'],
                        'style': block['style'],
                        'fingerprint': block['fingerprint']
                    })
            else:
                excluded_blocks.append({
                    'block_id': block['block_id'],
                    'text_preview': block['raw_text'][:50],
                    'reason': 'not_in_all_documents'
                })

        # Step 4: Determine header/footer handling
        # If headers differ, clear them
        header_fingerprints = [set(p['fingerprint'] for p in doc['header']) for doc in self.documents if doc['header']]
        clear_header = len(header_fingerprints) == 0 or len(set.intersection(*header_fingerprints) if header_fingerprints else set()) == 0

        # Build result
        result = {
            'version': '3.0',
            'created_at': datetime.now().isoformat(),
            'mode': self.mode,
            'analyzed_documents': len(self.documents),
            'source_files': [doc['doc_path'] for doc in self.documents],
            'reference_doc': reference_doc['doc_path'],
            'statistics': {
                'total_paragraphs_in_reference': reference_doc['paragraph_count'],
                'common_paragraphs': len(common_blocks),
                'excluded_paragraphs': len(excluded_blocks)
            },
            'common_fingerprints': list(common_fingerprints),
            'common_blocks': common_blocks,
            'excluded_blocks': excluded_blocks[:20],  # Limit for readability
            'clear_header': clear_header,
            'clear_footer': True,  # Usually safe to clear

            # For backward compatibility with existing code
            'sections': self._extract_sections_for_compatibility(common_blocks),
            'formatting': self._extract_formatting_from_reference()
        }

        return result

    def _extract_sections_for_compatibility(self, common_blocks: List[dict]) -> List[dict]:
        """Extract section info for backward compatibility with existing code."""
        sections = []
        known_headings = [
            "ANGABEN ZUR PERSON", "PERSONALIEN", "FAMILIENANAMNESE",
            "EIGENANAMNESE", "SOZIALANAMNESE", "BEFUND", "DIAGNOSE",
            "BEURTEILUNG", "ZUSAMMENFASSUNG", "EPIKRISE", "PROGNOSE",
            "MEDIKATION", "ANAMNESE", "VORGESCHICHTE"
        ]

        for i, block in enumerate(common_blocks):
            text_upper = block['text'].upper().strip()
            is_heading = any(h in text_upper for h in known_headings) or block['style'].startswith('Heading')

            if is_heading:
                sections.append({
                    'normalized_name': text_upper,
                    'display_name': block['text'].strip(),
                    'is_required': True,
                    'occurrence_count': len(self.documents),
                    'occurrence_percentage': 100.0,
                    'order': i
                })

        return sections

    def _extract_formatting_from_reference(self) -> dict:
        """Extract formatting info from reference document."""
        if not self.documents:
            return {'font_family': 'Times New Roman', 'font_size_pt': 12, 'line_spacing': 1.15}

        # Look at first few paragraphs to determine common formatting
        font_family = 'Times New Roman'
        font_size = 12

        for block in self.documents[0]['body'][:10]:
            for run in block.get('runs', []):
                if run.get('font_name'):
                    font_family = run['font_name']
                if run.get('font_size'):
                    try:
                        # Convert from Twips or EMUs to points
                        size = int(run['font_size'])
                        if size > 100:  # Likely EMUs or Twips
                            font_size = size // 12700  # EMUs to points
                        else:
                            font_size = size
                    except:
                        pass
                break
            break

        return {
            'font_family': font_family,
            'font_size_pt': font_size,
            'line_spacing': 1.15
        }

    def build_template_document(self, profile: dict, output_path: str) -> str:
        """Build the actual template DOCX file."""
        if not self.documents:
            raise ValueError("No documents analyzed")

        reference_path = profile['reference_doc']
        common_fingerprints = set(profile['common_fingerprints'])

        builder = TemplateBuilder(reference_path)
        return builder.build_template(
            keep_fingerprints=common_fingerprints,
            output_path=output_path,
            clear_header=profile.get('clear_header', True),
            clear_footer=profile.get('clear_footer', True),
            mode=self.mode
        )


def main():
    if len(sys.argv) < 3:
        print("Usage: python style_profile_analyzer.py <documents_json> <output_path>", file=sys.stderr)
        sys.exit(1)

    docs_input = sys.argv[1]
    output_path = sys.argv[2]

    # Parse input
    try:
        if os.path.isfile(docs_input):
            with open(docs_input, 'r', encoding='utf-8') as f:
                doc_paths = json.load(f)
        else:
            doc_paths = json.loads(docs_input)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}", file=sys.stderr)
        sys.exit(1)

    if not doc_paths:
        print("No documents provided", file=sys.stderr)
        sys.exit(1)

    # Analyze documents
    analyzer = StyleProfileAnalyzer(mode="COMMON")
    profile = analyzer.analyze_documents(doc_paths)

    # Save profile JSON
    try:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(profile, f, ensure_ascii=False, indent=2)
        print(f"Profile saved to: {output_path}", file=sys.stderr)
    except Exception as e:
        print(f"Error saving profile: {e}", file=sys.stderr)
        sys.exit(1)

    # Also build the template DOCX
    template_path = output_path.replace('.json', '_template.docx')
    try:
        analyzer.build_template_document(profile, template_path)
        print(f"Template DOCX saved to: {template_path}", file=sys.stderr)
    except Exception as e:
        print(f"Error building template: {e}", file=sys.stderr)

    # Output profile JSON to stdout for Rust to capture
    print(json.dumps(profile, ensure_ascii=False))


if __name__ == "__main__":
    main()
